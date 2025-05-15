
from docx import Document
import matplotlib.pyplot as plt
import seaborn as sns
from io import BytesIO

def generate_report(analysis_results):
    """Generate Word report from analysis results"""
    doc = Document()
    doc.add_heading('Log Analysis Report', 0)
    
    # Summary section
    doc.add_heading('Summary', level=1)
    total_logs = len(analysis_results)
    doc.add_paragraph(f"Total logs analyzed: {total_logs}")
    
    # Add visualizations
    doc.add_heading('Analysis Results', level=1)
    
    # Save report
    report_path = 'log-analyzer/reports/log_analysis_report.docx'
    doc.save(report_path)
    return report_path
